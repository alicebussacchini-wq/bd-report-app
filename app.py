import streamlit as st
import anthropic
import pdfplumber
import PyPDF2
import io
import os
import requests
import json
import base64
import time
import gspread
from pypdf import PdfReader, PdfWriter
from google.oauth2.service_account import Credentials
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── Configurazione pagina ─────────────────────────────────────────────────────

st.set_page_config(page_title="Taxi Report", page_icon="📊", layout="wide")

ARCHIVIO_PATH = r"C:\Users\1103540\bd-report-app\archivio"
os.makedirs(ARCHIVIO_PATH, exist_ok=True)

st.markdown("""
<style>
    .stApp { background-color: #1a1a1a; color: #f0f0f0; }
    .hl-header { display: flex; align-items: center; justify-content: space-between; padding: 20px 0; border-bottom: 3px solid #c8e04a; margin-bottom: 30px; }
    .hl-logo { height: 70px; }
    .hl-title { font-size: 22px; color: #c8e04a; font-weight: 600; text-align: right; }
    .hl-subtitle { font-size: 13px; color: #999; text-align: right; }
    .kpi-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(140px, 1fr)); gap: 12px; margin: 16px 0; }
    .kpi-card { background: #2a2a2a; border: 1px solid #c8e04a; border-radius: 10px; padding: 16px 12px; text-align: center; }
    .kpi-label { font-size: 11px; color: #999; text-transform: uppercase; letter-spacing: 1px; }
    .kpi-value { font-size: 20px; font-weight: 700; color: #c8e04a; margin-top: 6px; }
    .section-box { background: #242424; border-left: 4px solid #c8e04a; border-radius: 8px; padding: 20px; margin: 12px 0; }
    .section-title { font-size: 16px; font-weight: 700; color: #c8e04a; margin-bottom: 12px; }
    .section-text { font-size: 14px; color: #ddd; line-height: 1.7; }
    .ma-item { background: #2a2a2a; border-radius: 8px; padding: 14px; margin: 8px 0; border: 1px solid #444; }
    .ma-anno { background: #c8e04a; color: #1a1a1a; font-weight: 700; padding: 2px 10px; border-radius: 20px; font-size: 12px; display: inline-block; }
    .ma-tipo { color: #999; font-size: 12px; margin: 6px 0; text-transform: uppercase; }
    .ma-desc { color: #ddd; font-size: 14px; }
    .stButton > button { background: #c8e04a; color: #1a1a1a; font-weight: 700; border: none; padding: 12px 32px; border-radius: 8px; font-size: 16px; width: 100%; cursor: pointer; }
    .stButton > button:hover { background: #b5cc3a; }
    h1, h2, h3 { color: #f0f0f0; }
    .stExpander { background: #242424; border: 1px solid #444; border-radius: 8px; }
    .archivio-card { background: #242424; border: 1px solid #444; border-radius: 10px; padding: 18px; margin: 10px 0; }
    .archivio-nome { font-size: 18px; font-weight: 700; color: #c8e04a; }
    .archivio-data { font-size: 12px; color: #999; margin: 4px 0 12px 0; }
    .archivio-kpi { display: flex; gap: 12px; flex-wrap: wrap; }
    .archivio-kpi-item { background: #1a1a1a; border-radius: 6px; padding: 6px 12px; font-size: 12px; }
    .archivio-kpi-label { color: #999; }
    .archivio-kpi-val { color: #c8e04a; font-weight: 700; margin-left: 6px; }
    .pdf-badge { border-radius: 8px; padding: 10px 14px; margin: 8px 0; font-size: 13px; }
    .pdf-badge-ok { background: #1a2e1a; border: 1px solid #4caf50; color: #4caf50; }
    .pdf-badge-warn { background: #2e2a1a; border: 1px solid #ff9800; color: #ff9800; }
    .pdf-badge-scan { background: #1a1a2e; border: 1px solid #2196f3; color: #90caf9; }
</style>
""", unsafe_allow_html=True)

# ── Logo e header ─────────────────────────────────────────────────────────────

def get_logo_base64():
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo.jpg")
    if os.path.exists(logo_path):
        with open(logo_path, "rb") as f:
            return base64.b64encode(f.read()).decode()
    return None

logo_b64 = get_logo_base64()
logo_html = f'<img src="data:image/jpeg;base64,{logo_b64}" class="hl-logo">' if logo_b64 else "<span style='color:#c8e04a;font-size:24px;font-weight:700;'>Hogan Lovells</span>"
data_oggi = datetime.now().strftime("%d %B %Y")

st.markdown(f"""
<div class="hl-header">
    {logo_html}
    <div>
        <div class="hl-title">Taxi Report Generator</div>
        <div class="hl-subtitle">Generato il {data_oggi}</div>
    </div>
</div>
""", unsafe_allow_html=True)

# ── Costanti keyword per ricerca sezioni finanziarie ─────────────────────────

FINANCIAL_KEYWORDS = [
    ("stato patrimoniale",          10),
    ("conto economico",             10),
    ("rendiconto finanziario",      10),
    ("prospetto delle variazioni",   8),
    ("nota integrativa",             5),
    ("posizione finanziaria netta",  9),
    ("indebitamento finanziario",    9),
    ("pfn",                          6),
    ("ricavi delle vendite",         6),
    ("ricavi",                       4),
    ("ebitda",                       8),
    ("ebit",                         7),
    ("utile netto",                  6),
    ("perdita netto",                6),
    ("risultato netto",              6),
    ("risultato operativo",          6),
    ("ammortamenti",                 5),
    ("totale attivo",                6),
    ("totale passivo",               6),
    ("patrimonio netto",             7),
    ("capitale sociale",             4),
    ("debiti finanziari",            7),
    ("disponibilità liquide",        5),
    ("(migliaia di euro)",           8),
    ("(milioni di euro)",            8),
    ("in migliaia",                  7),
    ("in milioni",                   7),
    ("esercizio chiuso",             6),
    ("31 dicembre",                  5),
    ("31/12/",                       5),
]

SCORE_THRESHOLD = 12
CONTEXT_PAGES = 0

# ── Normalizzazione e validazione KPI ────────────────────────────────────────

# Fattori di conversione verso EURO INTERI
UNIT_FACTORS = {
    "euro": 1,
    "eur": 1,
    "€": 1,
    "migliaia di euro": 1_000,
    "migliaia": 1_000,
    "k€": 1_000,
    "milioni di euro": 1_000_000,
    "milioni": 1_000_000,
    "mln": 1_000_000,
    "mln€": 1_000_000,
    "miliardi di euro": 1_000_000_000,
    "miliardi": 1_000_000_000,
}

KPI_NUMERICI = [
    "ricavi", "ebitda", "ebit", "utile_netto", "ammortamenti",
    "totale_attivo", "patrimonio_netto", "cassa",
    "indebitamento_totale", "debito_bancario", "obbligazioni", "debito_netto",
]

def normalizza_kpi(report: dict) -> dict:
    """
    Normalizza tutti i KPI numerici a euro interi.
    Aggiunge campo '_normalizzato': True e '_unita_originale' al report.
    """
    fin = report.get("dati_finanziari", {})
    debito = report.get("struttura_debito", {})

    unita_raw = str(fin.get("unita", "migliaia di euro")).lower().strip()
    fattore = UNIT_FACTORS.get(unita_raw, None)

    # Prova match parziale se non trova esatto
    if fattore is None:
        for chiave, val in UNIT_FACTORS.items():
            if chiave in unita_raw:
                fattore = val
                break
    if fattore is None:
        fattore = 1_000  # default migliaia se non riconosciuto

    def converti(v):
        if v is None or v == "N/D" or v == "":
            return None
        try:
            return float(v) * fattore
        except (TypeError, ValueError):
            pass
        # Prova a estrarre numero e unità dalla stringa (es. "€ 1.2 miliardi")
        import re
        s = str(v).lower().strip()
        # Fattore dalla stringa stessa (per report vecchi senza campo "unita")
        str_factor = 1
        if "miliard" in s:
            str_factor = 1_000_000_000
        elif "milion" in s or " mln" in s or "mln " in s:
            str_factor = 1_000_000
        elif "migliaia" in s or " k" in s or "k€" in s:
            str_factor = 1_000
        # Se abbiamo già un fattore dalla stringa, non moltiplicare anche per fattore unità
        effective_factor = str_factor if str_factor > 1 else fattore
        nums = re.findall(r"-?[\d][\d.,]*", s)
        if nums:
            num_str = nums[0].replace(".", "").replace(",", ".")
            try:
                return float(num_str) * effective_factor
            except ValueError:
                pass
        return None

    # Normalizza dati finanziari
    fin_norm = {k: v for k, v in fin.items()}
    for campo in ["ricavi", "ebitda", "ebit", "utile_netto", "ammortamenti",
                  "totale_attivo", "patrimonio_netto", "cassa"]:
        fin_norm[campo] = converti(fin.get(campo))
    fin_norm["unita"] = "euro"
    fin_norm["_unita_originale"] = unita_raw
    fin_norm["_fattore"] = fattore

    # Normalizza struttura debito
    debito_norm = {k: v for k, v in debito.items()}
    for campo in ["indebitamento_totale", "debito_bancario", "obbligazioni", "debito_netto"]:
        debito_norm[campo] = converti(debito.get(campo))

    report["dati_finanziari"] = fin_norm
    report["struttura_debito"] = debito_norm
    report["_normalizzato"] = True
    return report


def valida_kpi(report: dict) -> list:
    """
    Esegue check di validazione sui KPI normalizzati.
    Restituisce una lista di warning (dict con 'livello', 'campo', 'messaggio').
    Livelli: 'error' (identità contabile violata), 'warning' (anomalia).
    """
    warnings = []
    fin = report.get("dati_finanziari", {})
    debito = report.get("struttura_debito", {})

    def get(d, k):
        v = d.get(k)
        return float(v) if v is not None else None

    ricavi      = get(fin, "ricavi")
    ebitda      = get(fin, "ebitda")
    ebit        = get(fin, "ebit")
    amm         = get(fin, "ammortamenti")
    utile       = get(fin, "utile_netto")
    attivo      = get(fin, "totale_attivo")
    pn          = get(fin, "patrimonio_netto")
    cassa       = get(fin, "cassa")
    deb_netto   = get(debito, "debito_netto")
    deb_tot     = get(debito, "indebitamento_totale")

    # ── Check identità contabili ─────────────────────────────────────────────

    # EBITDA = EBIT + Ammortamenti (tolleranza 5%)
    if ebitda is not None and ebit is not None and amm is not None:
        atteso = ebit + amm
        if atteso != 0 and abs(ebitda - atteso) / abs(atteso) > 0.05:
            warnings.append({
                "livello": "error",
                "campo": "EBITDA",
                "messaggio": f"EBITDA ({_fmt(ebitda)}) ≠ EBIT ({_fmt(ebit)}) + Amm. ({_fmt(amm)}) = {_fmt(atteso)}"
            })

    # Patrimonio Netto deve essere < Totale Attivo
    if pn is not None and attivo is not None and attivo > 0:
        if pn > attivo:
            warnings.append({
                "livello": "error",
                "campo": "Patrimonio Netto",
                "messaggio": f"Patrimonio netto ({_fmt(pn)}) > Totale attivo ({_fmt(attivo)}) — impossibile"
            })

    # ── Check anomalie di valore ─────────────────────────────────────────────

    # Margine EBITDA fuori range ragionevole (-50% / +80%)
    if ebitda is not None and ricavi is not None and ricavi > 0:
        margine = ebitda / ricavi
        if margine < -0.5 or margine > 0.8:
            warnings.append({
                "livello": "warning",
                "campo": "Margine EBITDA",
                "messaggio": f"Margine EBITDA {margine*100:.1f}% fuori range normale (-50%/+80%) — verificare unità"
            })

    # Cassa > Totale Attivo (impossibile)
    if cassa is not None and attivo is not None and attivo > 0:
        if cassa > attivo:
            warnings.append({
                "livello": "error",
                "campo": "Cassa",
                "messaggio": f"Cassa ({_fmt(cassa)}) > Totale attivo ({_fmt(attivo)}) — errore di estrazione"
            })

    # Ricavi negativi (quasi sempre errore)
    if ricavi is not None and ricavi < 0:
        warnings.append({
            "livello": "warning",
            "campo": "Ricavi",
            "messaggio": f"Ricavi negativi ({_fmt(ricavi)}) — verificare il dato estratto"
        })

    # Debito netto molto alto vs patrimonio netto (leva > 10x)
    if deb_netto is not None and pn is not None and pn > 0:
        leva = deb_netto / pn
        if leva > 10:
            warnings.append({
                "livello": "warning",
                "campo": "Leva finanziaria",
                "messaggio": f"Leva {leva:.1f}x molto elevata — possibile errore di estrazione"
            })

    return warnings


def confronta_anni(report_nuovo: dict, report_vecchio: dict) -> list:
    """
    Confronta due report dello stesso soggetto (anni diversi).
    Restituisce warning per variazioni YoY anomale (>200%).
    """
    warnings = []
    fin_n = report_nuovo.get("dati_finanziari", {})
    fin_v = report_vecchio.get("dati_finanziari", {})

    for campo in ["ricavi", "ebitda", "utile_netto", "totale_attivo", "patrimonio_netto"]:
        v_nuovo = fin_n.get(campo)
        v_vecchio = fin_v.get(campo)
        if v_nuovo is None or v_vecchio is None:
            continue
        try:
            v_nuovo, v_vecchio = float(v_nuovo), float(v_vecchio)
        except (TypeError, ValueError):
            continue
        if v_vecchio == 0:
            continue
        variazione = (v_nuovo - v_vecchio) / abs(v_vecchio)
        if abs(variazione) > 2.0:  # >200% YoY
            warnings.append({
                "livello": "warning",
                "campo": campo.replace("_", " ").title(),
                "messaggio": f"Variazione YoY {variazione*100:+.0f}% — possibile errore di estrazione o cambio unità"
            })
    return warnings


def _fmt(v):
    """Formatta un numero in euro in modo leggibile."""
    if v is None:
        return "N/D"
    try:
        v = float(v)
        if abs(v) >= 1_000_000_000:
            return f"€ {v/1_000_000_000:.2f}B"
        elif abs(v) >= 1_000_000:
            return f"€ {v/1_000_000:.1f}M"
        elif abs(v) >= 1_000:
            return f"€ {v/1_000:.0f}K"
        else:
            return f"€ {v:.0f}"
    except (TypeError, ValueError):
        return str(v)


def fmt_kpi(report: dict, campo: str, sezione: str = "dati_finanziari") -> str:
    """
    Formatta un KPI per la visualizzazione.
    Se il report è normalizzato usa _fmt, altrimenti usa il valore grezzo.
    """
    sezione_dict = report.get(sezione, {})
    v = sezione_dict.get(campo)
    if report.get("_normalizzato") and v is not None:
        return _fmt(v)
    if v is None:
        return "N/D"
    return str(v)


def render_pannello_validazione(warnings: list):
    """Renderizza il pannello di qualità estrazione con semaforo."""
    if not warnings:
        st.markdown(
            '<div style="background:#1a2e1a;border:1px solid #4caf50;border-radius:8px;'
            'padding:10px 14px;margin:8px 0;color:#4caf50;font-size:13px;">'
            '✅ <b>Validazione OK</b> — tutti i check superati</div>',
            unsafe_allow_html=True
        )
        return

    errori = [w for w in warnings if w["livello"] == "error"]
    avvisi = [w for w in warnings if w["livello"] == "warning"]

    if errori:
        colore, icona, label = "#c62828", "🔴", "Errori rilevati"
        bg = "#2e1a1a"
    else:
        colore, icona, label = "#e65100", "🟡", "Anomalie da verificare"
        bg = "#2e2a1a"

    items_html = "".join(
        f'<div style="margin:4px 0;font-size:12px;">'
        f'{"🔴" if w["livello"]=="error" else "🟡"} '
        f'<b>{w["campo"]}</b>: {w["messaggio"]}</div>'
        for w in warnings
    )
    st.markdown(
        f'<div style="background:{bg};border:1px solid {colore};border-radius:8px;'
        f'padding:10px 14px;margin:8px 0;">'
        f'<div style="color:{colore};font-weight:700;margin-bottom:6px;">{icona} {label}</div>'
        f'{items_html}</div>',
        unsafe_allow_html=True
    )


# ── Funzioni PDF intelligente ─────────────────────────────────────────────────

def _score_page(text: str) -> int:
    text_lower = text.lower()
    score = 0
    for keyword, weight in FINANCIAL_KEYWORDS:
        if keyword in text_lower:
            score += weight
    return score


def analizza_pdf(pdf_bytes: bytes) -> dict:
    """
    Analizza un PDF e trova le pagine con dati finanziari.
    Restituisce info sull'analisi e la modalità di estrazione.
    """
    result = {
        "mode": "text",
        "total_pages": 0,
        "selected_pages": [],
        "page_scores": [],
        "avg_chars": 0.0,
        "sections_found": [],
    }

    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            total = len(pdf.pages)
            result["total_pages"] = total

            texts, scores, char_counts = [], [], []
            for page in pdf.pages:
                t = page.extract_text() or ""
                texts.append(t)
                scores.append(_score_page(t))
                char_counts.append(len(t))

            result["page_scores"] = scores
            avg_chars = sum(char_counts) / total if total > 0 else 0
            result["avg_chars"] = avg_chars

            # PDF scansionato: meno di 100 caratteri per pagina in media
            if avg_chars < 100:
                result["mode"] = "scanned"
                result["selected_pages"] = list(range(total))
                return result

            # Selezione pagine rilevanti + contesto
            relevant = set()
            for i, score in enumerate(scores):
                if score >= SCORE_THRESHOLD:
                    for j in range(
                        max(0, i - CONTEXT_PAGES),
                        min(total, i + CONTEXT_PAGES + 1)
                    ):
                        relevant.add(j)

            # Fallback: nessuna sezione trovata → prendiamo la metà centrale
            if not relevant:
                start = max(0, total // 4)
                end = min(total, 3 * total // 4)
                relevant = set(range(start, end))
                result["mode"] = "fallback"

            # Cap massimo 25 pagine per rispettare i rate limit API
            selected_sorted = sorted(relevant)
            if len(selected_sorted) > 25:
                # Tieni le pagine con punteggio più alto
                scored = sorted(selected_sorted, key=lambda i: scores[i], reverse=True)
                selected_sorted = sorted(scored[:25])
            result["selected_pages"] = selected_sorted

            # Sezioni identificate
            all_text = " ".join(texts).lower()
            sections = []
            for kw, label in [
                ("stato patrimoniale", "Stato Patrimoniale"),
                ("conto economico", "Conto Economico"),
                ("rendiconto finanziario", "Rendiconto Finanziario"),
                ("posizione finanziaria netta", "PFN"),
                ("nota integrativa", "Nota Integrativa"),
            ]:
                if kw in all_text:
                    sections.append(label)
            result["sections_found"] = sections

    except Exception as e:
        # Se pdfplumber fallisce, trattalo come scansionato
        result["mode"] = "scanned"
        result["error"] = str(e)
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
            result["total_pages"] = len(reader.pages)
            result["selected_pages"] = list(range(len(reader.pages)))
        except:
            pass

    return result


def costruisci_pdf_chirurgico(pdf_bytes: bytes, page_indices: list) -> bytes:
    """Costruisce un PDF con solo le pagine selezionate (0-indexed)."""
    reader = PdfReader(io.BytesIO(pdf_bytes))
    writer = PdfWriter()
    for i in page_indices:
        if 0 <= i < len(reader.pages):
            writer.add_page(reader.pages[i])
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def prepara_bilancio(pdf_bytes: bytes) -> tuple:
    """
    Entry point principale per preparare un bilancio per Claude.
    Restituisce (pdf_da_inviare_bytes, info_dict).
    """
    info = analizza_pdf(pdf_bytes)

    if info["mode"] == "scanned":
        return pdf_bytes, info

    surgical = costruisci_pdf_chirurgico(pdf_bytes, info["selected_pages"])
    info["pagine_inviate"] = len(info["selected_pages"])
    info["riduzione_pct"] = round(
        (1 - len(surgical) / len(pdf_bytes)) * 100, 1
    ) if len(pdf_bytes) > 0 else 0

    return surgical, info

# ── Funzioni Google Sheets ────────────────────────────────────────────────────

def estrai_testo_url(url):
    try:
        return requests.get(url, timeout=10).text[:5000]
    except:
        return "Impossibile recuperare il contenuto dell'URL."

def get_sheet():
    credenziali = {
        "type": st.secrets["gcp_service_account"]["type"],
        "project_id": st.secrets["gcp_service_account"]["project_id"],
        "private_key_id": st.secrets["gcp_service_account"]["private_key_id"],
        "private_key": st.secrets["gcp_service_account"]["private_key"],
        "client_email": st.secrets["gcp_service_account"]["client_email"],
        "client_id": st.secrets["gcp_service_account"]["client_id"],
        "auth_uri": st.secrets["gcp_service_account"]["auth_uri"],
        "token_uri": st.secrets["gcp_service_account"]["token_uri"],
        "auth_provider_x509_cert_url": st.secrets["gcp_service_account"]["auth_provider_x509_cert_url"],
        "client_x509_cert_url": st.secrets["gcp_service_account"]["client_x509_cert_url"]
    }
    scopes = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    creds = Credentials.from_service_account_info(credenziali, scopes=scopes)
    client = gspread.authorize(creds)
    sheet = client.open("Taxi Report Archivio").sheet1
    return sheet

def salva_report(nome_azienda, report_json, documenti_files):
    try:
        sheet = get_sheet()
        timestamp = datetime.now().strftime("%d/%m/%Y %H:%M")
        sheet.append_row([
            timestamp,
            nome_azienda,
            json.dumps(report_json, ensure_ascii=False)
        ])
    except Exception as e:
        st.warning(f"Errore salvataggio archivio: {e}")

def carica_archivio():
    try:
        sheet = get_sheet()
        righe = sheet.get_all_values()
        reports = []
        for riga in reversed(righe):
            if len(riga) >= 3 and riga[2]:
                try:
                    report = json.loads(riga[2])
                    # Normalizza sempre i KPI al caricamento,
                    # così i report vecchi (pre-M2) vengono convertiti on-the-fly
                    if not report.get("_normalizzato"):
                        report = normalizza_kpi(report)
                    reports.append({
                        "data": riga[0],
                        "nome": riga[1],
                        "report": report,
                        "riga": righe.index(riga) + 1
                    })
                except:
                    pass
        return reports
    except Exception as e:
        st.warning(f"Errore caricamento archivio: {e}")
        return []


# ── Esportazione Excel ────────────────────────────────────────────────────────

# Colori HL
HL_VERDE    = "C8E04A"
HL_NERO     = "1A1A1A"
HL_GRIGIO_S = "2A2A2A"
HL_GRIGIO_C = "444444"
HL_BIANCO   = "F0F0F0"

def _cell_style(ws, row, col, value, bold=False, bg=None, fg="F0F0F0",
                align="left", border=False, num_format=None):
    cell = ws.cell(row=row, column=col, value=value)
    cell.font = Font(name="Calibri", size=10, bold=bold,
                     color=fg if fg else "000000")
    if bg:
        cell.fill = PatternFill("solid", fgColor=bg)
    cell.alignment = Alignment(horizontal=align, vertical="center",
                               wrap_text=True)
    if border:
        side = Side(style="thin", color=HL_GRIGIO_C)
        cell.border = Border(left=side, right=side, top=side, bottom=side)
    if num_format:
        cell.number_format = num_format
    return cell


def genera_excel(reports: list) -> bytes:
    """
    Genera un file Excel con un foglio per ogni report
    più un foglio KPI Summary comparativo se ci sono più report.
    """
    wb = openpyxl.Workbook()
    wb.remove(wb.active)  # rimuovi foglio vuoto default

    # ── Foglio KPI Summary (sempre presente) ─────────────────────────────────
    ws_sum = wb.create_sheet("KPI Summary")
    ws_sum.sheet_view.showGridLines = False
    ws_sum.tab_color = HL_VERDE

    # Header riga 1
    ws_sum.row_dimensions[1].height = 30
    titolo_cell = ws_sum.cell(row=1, column=1, value="TAXI REPORT — KPI Summary")
    titolo_cell.font = Font(name="Calibri", size=14, bold=True, color=HL_VERDE)
    titolo_cell.fill = PatternFill("solid", fgColor=HL_NERO)
    ws_sum.merge_cells(start_row=1, start_column=1,
                       end_row=1, end_column=2 + len(reports))

    # Sottotitolo
    ws_sum.row_dimensions[2].height = 16
    sub = ws_sum.cell(row=2, column=1,
                      value=f"Generato il {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    sub.font = Font(name="Calibri", size=9, color="999999")
    sub.fill = PatternFill("solid", fgColor=HL_NERO)
    ws_sum.merge_cells(start_row=2, start_column=1,
                       end_row=2, end_column=2 + len(reports))

    # Intestazioni colonne: KPI | Azienda1 (anno) | Azienda2 (anno) ...
    ws_sum.row_dimensions[4].height = 28
    _cell_style(ws_sum, 4, 1, "KPI", bold=True, bg=HL_NERO, fg=HL_VERDE,
                align="left", border=True)
    _cell_style(ws_sum, 4, 2, "Unità", bold=True, bg=HL_NERO, fg=HL_VERDE,
                align="center", border=True)
    for j, r in enumerate(reports):
        label = f"{r['nome']} ({r['anno']})"
        _cell_style(ws_sum, 4, 3 + j, label, bold=True, bg=HL_NERO,
                    fg=HL_VERDE, align="center", border=True)

    # Righe KPI
    kpi_rows = [
        ("Ricavi",           "dati_finanziari", "ricavi",            "€ #,##0"),
        ("EBITDA",           "dati_finanziari", "ebitda",            "€ #,##0"),
        ("Margine EBITDA",   None,              "_margine_ebitda",   "0.0%"),
        ("EBIT",             "dati_finanziari", "ebit",              "€ #,##0"),
        ("Utile Netto",      "dati_finanziari", "utile_netto",       "€ #,##0"),
        ("Ammortamenti",     "dati_finanziari", "ammortamenti",      "€ #,##0"),
        ("Totale Attivo",    "dati_finanziari", "totale_attivo",     "€ #,##0"),
        ("Patrimonio Netto", "dati_finanziari", "patrimonio_netto",  "€ #,##0"),
        ("Cassa",            "dati_finanziari", "cassa",             "€ #,##0"),
        ("Debito Netto",     "struttura_debito","debito_netto",      "€ #,##0"),
        ("Indebitamento Tot","struttura_debito","indebitamento_totale","€ #,##0"),
        ("Leva (x)",         "struttura_debito","leva_finanziaria",  "0.0"),
    ]

    for i, (label, sezione, campo, fmt) in enumerate(kpi_rows):
        row_idx = 5 + i
        ws_sum.row_dimensions[row_idx].height = 20
        bg = HL_GRIGIO_S if i % 2 == 0 else HL_NERO
        _cell_style(ws_sum, row_idx, 1, label, bg=bg, fg=HL_BIANCO,
                    border=True)
        _cell_style(ws_sum, row_idx, 2, "EUR" if "€" in fmt else "x" if "0.0" == fmt else "%",
                    bg=bg, fg="999999", align="center", border=True)

        for j, r in enumerate(reports):
            report = r["report"]
            val = None
            if campo == "_margine_ebitda":
                ric = report.get("dati_finanziari", {}).get("ricavi")
                ebt = report.get("dati_finanziari", {}).get("ebitda")
                if ric and ebt:
                    try:
                        val = float(ebt) / float(ric)
                    except:
                        pass
            elif sezione:
                raw = report.get(sezione, {}).get(campo)
                if raw is not None:
                    try:
                        val = float(raw)
                    except:
                        pass

            cell = _cell_style(ws_sum, row_idx, 3 + j, val, bg=bg,
                               fg=HL_BIANCO, align="right", border=True,
                               num_format=fmt)

    # Warning validazione se presenti
    warn_row = 5 + len(kpi_rows) + 2
    has_warnings = any(r["report"].get("_validation_warnings") for r in reports)
    if has_warnings:
        w_cell = ws_sum.cell(row=warn_row, column=1,
                             value="⚠️ Note validazione")
        w_cell.font = Font(name="Calibri", size=10, bold=True, color="E65100")
        w_cell.fill = PatternFill("solid", fgColor="2E2A1A")
        ws_sum.merge_cells(start_row=warn_row, start_column=1,
                           end_row=warn_row, end_column=2 + len(reports))
        for r in reports:
            for w in r["report"].get("_validation_warnings", []):
                warn_row += 1
                msg = f"{r['nome']} ({r['anno']}) — {w['campo']}: {w['messaggio']}"
                wc = ws_sum.cell(row=warn_row, column=1, value=msg)
                wc.font = Font(name="Calibri", size=9,
                               color="C62828" if w["livello"] == "error" else "E65100")
                wc.fill = PatternFill("solid", fgColor=HL_NERO)
                ws_sum.merge_cells(start_row=warn_row, start_column=1,
                                   end_row=warn_row, end_column=2 + len(reports))

    # Larghezze colonne
    ws_sum.column_dimensions["A"].width = 22
    ws_sum.column_dimensions["B"].width = 8
    for j in range(len(reports)):
        ws_sum.column_dimensions[get_column_letter(3 + j)].width = 22

    # ── Un foglio per ogni report (dati completi) ─────────────────────────────
    for r in reports:
        nome_sheet = f"{r['nome'][:20]} {r['anno']}"
        ws = wb.create_sheet(nome_sheet)
        ws.sheet_view.showGridLines = False

        report = r["report"]
        fin    = report.get("dati_finanziari", {})
        debito = report.get("struttura_debito", {})

        # Intestazione
        ws.row_dimensions[1].height = 30
        h = ws.cell(row=1, column=1,
                    value=f"{r['nome']} — Bilancio {r['anno']}")
        h.font = Font(name="Calibri", size=13, bold=True, color=HL_VERDE)
        h.fill = PatternFill("solid", fgColor=HL_NERO)
        ws.merge_cells("A1:C1")

        # Sezione Dati Finanziari
        _cell_style(ws, 3, 1, "DATI FINANZIARI", bold=True, bg=HL_GRIGIO_S,
                    fg=HL_VERDE, border=True)
        _cell_style(ws, 3, 2, "Valore (EUR)", bold=True, bg=HL_GRIGIO_S,
                    fg=HL_VERDE, align="right", border=True)
        _cell_style(ws, 3, 3, "Note", bold=True, bg=HL_GRIGIO_S,
                    fg=HL_VERDE, border=True)

        fin_campi = [
            ("Ricavi",            fin.get("ricavi")),
            ("EBITDA",            fin.get("ebitda")),
            ("EBIT",              fin.get("ebit")),
            ("Utile Netto",       fin.get("utile_netto")),
            ("Ammortamenti",      fin.get("ammortamenti")),
            ("Totale Attivo",     fin.get("totale_attivo")),
            ("Patrimonio Netto",  fin.get("patrimonio_netto")),
            ("Cassa",             fin.get("cassa")),
        ]
        for i, (label, val) in enumerate(fin_campi):
            row = 4 + i
            bg = HL_GRIGIO_S if i % 2 == 0 else HL_NERO
            ws.row_dimensions[row].height = 18
            _cell_style(ws, row, 1, label, bg=bg, fg=HL_BIANCO, border=True)
            num = None
            if val is not None:
                try:
                    num = float(val)
                except:
                    pass
            _cell_style(ws, row, 2, num, bg=bg, fg=HL_BIANCO,
                        align="right", border=True, num_format="€ #,##0")
            nota = ""
            if label == "Unità originale":
                nota = fin.get("_unita_originale", "")
            _cell_style(ws, row, 3, nota, bg=bg, fg="999999", border=True)

        # Sezione Struttura Debito
        deb_start = 4 + len(fin_campi) + 2
        _cell_style(ws, deb_start, 1, "STRUTTURA DEBITO", bold=True,
                    bg=HL_GRIGIO_S, fg=HL_VERDE, border=True)
        _cell_style(ws, deb_start, 2, "Valore (EUR)", bold=True,
                    bg=HL_GRIGIO_S, fg=HL_VERDE, align="right", border=True)
        _cell_style(ws, deb_start, 3, "Note", bold=True, bg=HL_GRIGIO_S,
                    fg=HL_VERDE, border=True)

        deb_campi = [
            ("Indebitamento Totale", debito.get("indebitamento_totale")),
            ("Debito Bancario",      debito.get("debito_bancario")),
            ("Obbligazioni",         debito.get("obbligazioni")),
            ("Debito Netto",         debito.get("debito_netto")),
            ("Leva Finanziaria",     debito.get("leva_finanziaria")),
        ]
        for i, (label, val) in enumerate(deb_campi):
            row = deb_start + 1 + i
            bg = HL_GRIGIO_S if i % 2 == 0 else HL_NERO
            ws.row_dimensions[row].height = 18
            _cell_style(ws, row, 1, label, bg=bg, fg=HL_BIANCO, border=True)
            num = None
            if val is not None:
                try:
                    num = float(val)
                except:
                    pass
            fmt = "0.0" if label == "Leva Finanziaria" else "€ #,##0"
            _cell_style(ws, row, 2, num, bg=bg, fg=HL_BIANCO,
                        align="right", border=True, num_format=fmt)
            _cell_style(ws, row, 3, "", bg=bg, fg="999999", border=True)

        # Sezione testi
        testi_start = deb_start + len(deb_campi) + 2
        for label, campo in [
            ("OVERVIEW", "overview"),
            ("CORE BUSINESS", "core_business"),
            ("MERCATI", "mercati"),
            ("NOTE AGGIUNTIVE", "note_aggiuntive"),
        ]:
            val = report.get(campo, "N/D") or "N/D"
            _cell_style(ws, testi_start, 1, label, bold=True,
                        bg=HL_GRIGIO_S, fg=HL_VERDE, border=True)
            ws.merge_cells(start_row=testi_start, start_column=1,
                           end_row=testi_start, end_column=3)
            testi_start += 1
            tc = ws.cell(row=testi_start, column=1, value=val)
            tc.font = Font(name="Calibri", size=9, color=HL_BIANCO)
            tc.fill = PatternFill("solid", fgColor=HL_NERO)
            tc.alignment = Alignment(wrap_text=True, vertical="top")
            ws.row_dimensions[testi_start].height = 60
            ws.merge_cells(start_row=testi_start, start_column=1,
                           end_row=testi_start, end_column=3)
            testi_start += 2

        # Larghezze
        ws.column_dimensions["A"].width = 25
        ws.column_dimensions["B"].width = 20
        ws.column_dimensions["C"].width = 35

    # Output bytes
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Esportazione Word ─────────────────────────────────────────────────────────

def genera_word(reports: list) -> bytes:
    """
    Genera un documento Word con un report per azienda.
    Stile HL: header con logo/titolo, tabelle KPI, sezioni testo.
    """
    doc = Document()

    # Imposta margini
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Stili base
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(10)

    def add_heading(text, level=1, color_hex="C8E04A"):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.bold = True
        run.font.size = Pt(14 if level == 1 else 11)
        run.font.color.rgb = RGBColor.from_string(color_hex)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def add_kpi_table(doc, kpi_list):
        """Aggiunge una tabella KPI a due colonne (label | valore)."""
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Larghezze colonne
        for col_idx, width in enumerate([Cm(6), Cm(10)]):
            for cell in table.columns[col_idx].cells:
                cell.width = width

        for label, value in kpi_list:
            row = table.add_row()
            # Label
            lc = row.cells[0]
            lc.text = label
            lc.paragraphs[0].runs[0].font.bold = True
            lc.paragraphs[0].runs[0].font.size = Pt(9)
            lc.paragraphs[0].runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # Valore
            vc = row.cells[1]
            vc.text = str(value) if value is not None else "N/D"
            vc.paragraphs[0].runs[0].font.size = Pt(10)
            vc.paragraphs[0].runs[0].font.bold = True
        return table

    for idx, r in enumerate(reports):
        report = r["report"]
        fin    = report.get("dati_finanziari", {})
        debito = report.get("struttura_debito", {})

        # ── Copertina sezione ────────────────────────────────────────────────
        if idx > 0:
            doc.add_page_break()

        # Titolo azienda
        p_title = doc.add_paragraph()
        run = p_title.add_run(r["nome"].upper())
        run.font.name   = "Calibri"
        run.font.bold   = True
        run.font.size   = Pt(20)
        run.font.color.rgb = RGBColor.from_string("C8E04A")
        p_title.paragraph_format.space_after = Pt(2)

        p_sub = doc.add_paragraph()
        run2 = p_sub.add_run(f"Bilancio {r['anno']}  |  Generato il {datetime.now().strftime('%d/%m/%Y')}")
        run2.font.name  = "Calibri"
        run2.font.size  = Pt(9)
        run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        p_sub.paragraph_format.space_after = Pt(16)

        # Linea separatore
        p_hr = doc.add_paragraph("─" * 80)
        p_hr.runs[0].font.color.rgb = RGBColor.from_string("C8E04A")
        p_hr.runs[0].font.size = Pt(7)
        p_hr.paragraph_format.space_after = Pt(12)

        # ── Warning validazione ───────────────────────────────────────────────
        warnings = report.get("_validation_warnings", [])
        if warnings:
            add_heading("⚠️ Note di validazione", level=2, color_hex="E65100")
            for w in warnings:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"{w['campo']}: {w['messaggio']}")
                run.font.size = Pt(9)
                run.font.color.rgb = (RGBColor(0xC6, 0x28, 0x28)
                                      if w["livello"] == "error"
                                      else RGBColor(0xE6, 0x51, 0x00))

        # ── Dati finanziari ───────────────────────────────────────────────────
        add_heading("Dati Finanziari", level=2)
        kpi_fin = [
            ("Ricavi",            _fmt(fin.get("ricavi"))),
            ("EBITDA",            _fmt(fin.get("ebitda"))),
            ("Margine EBITDA",    f"{float(fin['ebitda'])/float(fin['ricavi'])*100:.1f}%"
                                  if fin.get("ebitda") and fin.get("ricavi") else "N/D"),
            ("EBIT",              _fmt(fin.get("ebit"))),
            ("Utile Netto",       _fmt(fin.get("utile_netto"))),
            ("Totale Attivo",     _fmt(fin.get("totale_attivo"))),
            ("Patrimonio Netto",  _fmt(fin.get("patrimonio_netto"))),
            ("Cassa",             _fmt(fin.get("cassa"))),
        ]
        add_kpi_table(doc, kpi_fin)

        # ── Struttura debito ──────────────────────────────────────────────────
        add_heading("Struttura del Debito", level=2)
        kpi_deb = [
            ("Indebitamento Totale", _fmt(debito.get("indebitamento_totale"))),
            ("Debito Bancario",      _fmt(debito.get("debito_bancario"))),
            ("Obbligazioni",         _fmt(debito.get("obbligazioni"))),
            ("Debito Netto",         _fmt(debito.get("debito_netto"))),
            ("Leva Finanziaria",     str(debito.get("leva_finanziaria", "N/D"))),
            ("Scadenze",             debito.get("scadenze_principali", "N/D") or "N/D"),
        ]
        add_kpi_table(doc, kpi_deb)

        # ── Sezioni testuali ──────────────────────────────────────────────────
        for label, campo in [
            ("Overview",       "overview"),
            ("Core Business",  "core_business"),
            ("Mercati",        "mercati"),
            ("Ownership",      None),
            ("Operazioni M&A", None),
            ("Note Aggiuntive","note_aggiuntive"),
        ]:
            if campo:
                testo = report.get(campo, "N/D") or "N/D"
            elif label == "Ownership":
                own = report.get("ownership", {})
                if isinstance(own, dict):
                    testo = (f"Azionista principale: {own.get('azionista_principale','N/D')} "
                             f"({own.get('quota_principale','N/D')}). "
                             f"{own.get('struttura_controllo','')}")
                else:
                    testo = str(own)
            else:  # Operazioni M&A
                ops = report.get("operazioni_ma", [])
                testo = "; ".join(
                    f"{o.get('anno','')} — {o.get('tipo','')} — {o.get('descrizione','')}"
                    for o in ops
                    if o.get("descrizione","N/D") != "N/D"
                ) or "N/D"

            if testo and testo != "N/D":
                add_heading(label, level=2)
                p = doc.add_paragraph(testo)
                p.runs[0].font.size = Pt(10)
                p.paragraph_format.space_after = Pt(6)

    # ── Footer nota metodologica ──────────────────────────────────────────────
    doc.add_page_break()
    add_heading("Note Metodologiche", level=1, color_hex="999999")
    p = doc.add_paragraph(
        "I dati finanziari sono stati estratti automaticamente dal bilancio tramite "
        "analisi AI (Claude). I KPI sono stati normalizzati in euro. "
        "Si raccomanda di verificare i valori rispetto al documento originale "
        "prima di qualsiasi utilizzo professionale. "
        f"Report generato il {datetime.now().strftime('%d/%m/%Y alle %H:%M')}."
    )
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# BENCHMARK DI SETTORE (statico)
# ══════════════════════════════════════════════════════════════════════════════

BENCHMARK_SETTORE = {
    "Farmaceutico / Healthcare": {
        "margine_ebitda": (0.18, 0.28),
        "margine_utile":  (0.08, 0.15),
        "leva":           (1.0,  3.0),
    },
    "Manifatturiero": {
        "margine_ebitda": (0.08, 0.16),
        "margine_utile":  (0.03, 0.08),
        "leva":           (1.5,  4.0),
    },
    "Retail / Grande Distribuzione": {
        "margine_ebitda": (0.04, 0.10),
        "margine_utile":  (0.01, 0.04),
        "leva":           (2.0,  5.0),
    },
    "Technology / Software": {
        "margine_ebitda": (0.15, 0.35),
        "margine_utile":  (0.08, 0.20),
        "leva":           (0.0,  2.0),
    },
    "Energia / Utilities": {
        "margine_ebitda": (0.20, 0.40),
        "margine_utile":  (0.05, 0.12),
        "leva":           (3.0,  6.0),
    },
    "Costruzioni / Real Estate": {
        "margine_ebitda": (0.08, 0.18),
        "margine_utile":  (0.03, 0.08),
        "leva":           (2.0,  6.0),
    },
    "Alimentare / Bevande": {
        "margine_ebitda": (0.10, 0.20),
        "margine_utile":  (0.04, 0.10),
        "leva":           (1.0,  3.5),
    },
}

def _benchmark_cell(valore, bm_range):
    """Restituisce (testo, colore) per una cella benchmark."""
    if valore is None or bm_range is None:
        return "—", "#666"
    lo, hi = bm_range
    if lo <= valore <= hi:
        return "✅ In range", "#4caf50"
    elif valore < lo:
        diff = (lo - valore) / lo * 100
        return f"🔴 -{diff:.0f}% vs mediana", "#ef5350"
    else:
        diff = (valore - hi) / hi * 100
        return f"🟡 +{diff:.0f}% vs mediana", "#ffa726"

# ── Navigazione ───────────────────────────────────────────────────────────────

if "pagina" not in st.session_state:
    st.session_state["pagina"] = "genera"

col_nav1, col_nav2, col_nav3, col_nav4 = st.columns(4)
with col_nav1:
    if st.button("➕ Genera nuovo report", use_container_width=True):
        st.session_state["pagina"] = "genera"
with col_nav2:
    if st.button("🗂️ Archivio report", use_container_width=True):
        st.session_state["pagina"] = "archivio"
with col_nav3:
    if st.button("📊 Confronta aziende", use_container_width=True):
        st.session_state["pagina"] = "archivio"  # apre archivio sul tab confronta
with col_nav4:
    if st.button("📄 Carica report esistente", use_container_width=True):
        st.session_state["pagina"] = "carica_esistente"

st.markdown("---")

# ══════════════════════════════════════════════════════════════════════════════
# PAGINA: GENERA REPORT
# ══════════════════════════════════════════════════════════════════════════════

if st.session_state["pagina"] == "genera":

    lingua = st.radio("🌐 Lingua del report", ["Italiano", "English"], horizontal=True)

    # ── SEZIONE BILANCI (multi-upload) ────────────────────────────────────────

    st.markdown("### 📂 Carica i bilanci")
    st.caption(
        "Puoi caricare fino a 5 bilanci PDF completi contemporaneamente. "
        "Il sistema trova automaticamente le pagine finanziarie rilevanti — "
        "non serve estrarre le pagine manualmente."
    )

    bilanci_files = st.file_uploader(
        "Seleziona uno o più bilanci PDF",
        type=["pdf"],
        accept_multiple_files=True,
        key="multi_bilanci",
        label_visibility="collapsed",
    )

    # Form metadati + analisi per ogni bilancio caricato
    bilanci_pronti = []   # lista di dict pronti per Claude

    if bilanci_files:
        if len(bilanci_files) > 5:
            st.error("Massimo 5 bilanci alla volta.")
            bilanci_files = bilanci_files[:5]

        st.markdown("#### Conferma i dettagli")

        for idx, f in enumerate(bilanci_files):
            with st.container():
                st.markdown(f"**📄 {f.name}**")
                col_m1, col_m2 = st.columns([3, 1])
                with col_m1:
                    rs = st.text_input(
                        "Ragione sociale",
                        key=f"rs_{idx}",
                        placeholder="es. Zambon S.p.A.",
                    )
                with col_m2:
                    anno = st.selectbox(
                        "Anno bilancio",
                        options=[str(y) for y in range(2025, 2017, -1)],
                        key=f"anno_{idx}",
                    )

                # Analisi pagine (cached in session_state per evitare ri-analisi)
                cache_key = f"pdf_analysis_{f.name}_{f.size}"
                if cache_key not in st.session_state:
                    with st.spinner(f"Analisi pagine di {f.name}..."):
                        raw_bytes = f.read()
                        surgical_bytes, info = prepara_bilancio(raw_bytes)
                        st.session_state[cache_key] = {
                            "surgical": surgical_bytes,
                            "info": info,
                        }
                else:
                    surgical_bytes = st.session_state[cache_key]["surgical"]
                    info = st.session_state[cache_key]["info"]

                # Badge risultato analisi
                mode = info.get("mode", "text")
                total = info.get("total_pages", 0)
                selected = len(info.get("selected_pages", []))
                sections = ", ".join(info.get("sections_found", [])) or "—"
                riduzione = info.get("riduzione_pct", 0)

                if mode == "scanned":
                    st.markdown(
                        f'<div class="pdf-badge pdf-badge-scan">'
                        f'📷 PDF scansionato — {total} pagine totali inviate a Claude direttamente.'
                        f'</div>',
                        unsafe_allow_html=True,
                    )
                elif mode == "fallback":
                    st.markdown(
                        f'<div class="pdf-badge pdf-badge-warn">'
                        f'⚠️ Layout atipico — invio {selected} pagine centrali su {total} totali.'
                        f'</div>',
                        unsafe_allow_html=True,
                    )
                else:
                    st.markdown(
                        f'<div class="pdf-badge pdf-badge-ok">'
                        f'✅ {selected} pagine rilevanti su {total} totali '
                        f'({riduzione}% riduzione) — Sezioni: {sections}'
                        f'</div>',
                        unsafe_allow_html=True,
                    )

                # Pannello debug opzionale
                with st.expander("🔍 Dettaglio analisi pagine", expanded=False):
                    scores = info.get("page_scores", [])
                    selected_set = set(info.get("selected_pages", []))
                    if scores:
                        st.caption(f"Soglia punteggio: {SCORE_THRESHOLD} | "
                                   f"Caratteri medi per pagina: {info.get('avg_chars', 0):.0f}")
                        cols_dbg = st.columns(10)
                        for i, score in enumerate(scores[:50]):
                            col = cols_dbg[i % 10]
                            color = "#c8e04a" if i in selected_set else "#555"
                            col.markdown(
                                f"<div style='text-align:center;font-size:10px;"
                                f"color:{color};border:1px solid {color};"
                                f"border-radius:4px;padding:3px;margin:2px;'>"
                                f"<b>{i+1}</b><br>{score}</div>",
                                unsafe_allow_html=True,
                            )
                        if len(scores) > 50:
                            st.caption(f"... e altre {len(scores)-50} pagine.")

                if rs.strip():
                    bilanci_pronti.append({
                        "ragione_sociale": rs.strip(),
                        "anno": anno,
                        "pdf_b64": base64.b64encode(surgical_bytes).decode(),
                        "filename": f.name,
                        "info": info,
                    })
                else:
                    st.warning("⬆️ Inserisci la ragione sociale per includere questo bilancio.")

                st.markdown("---")

    # ── ALTRI DOCUMENTI (invariati) ───────────────────────────────────────────

    st.markdown("### 📎 Altri documenti (opzionale)")
    col1, col2, col3 = st.columns(3)

    with col1:
        st.caption("📈 Export Mergermarket")
        mergermarket = st.file_uploader(
            "Mergermarket", type=["pdf", "csv"],
            key="merger", label_visibility="collapsed"
        )

    with col2:
        st.caption("🌐 URL Sito / Press Release")
        url_azienda = st.text_input(
            "URL", placeholder="https://...", label_visibility="collapsed"
        )

    with col3:
        st.caption("🏛️ Visura Camerale")
        visura = st.file_uploader(
            "Visura PDF", type=["pdf"],
            key="visura", label_visibility="collapsed"
        )

    # Raccolta testi documenti supplementari
    testi_supplementari = {}
    documenti_binari = {}

    if mergermarket:
        contenuto = mergermarket.read()
        if mergermarket.name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(io.BytesIO(contenuto))
            testo = "".join(p.extract_text() or "" for p in reader.pages)
        else:
            testo = contenuto.decode("utf-8")
        testi_supplementari["Mergermarket"] = testo
        documenti_binari[mergermarket.name] = contenuto
        st.success("✅ Mergermarket caricato")

    if url_azienda:
        testi_supplementari["Sito Aziendale"] = estrai_testo_url(url_azienda)
        st.success("✅ URL acquisito")

    if visura:
        contenuto = visura.read()
        try:
            client_temp = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))
            resp_visura = client_temp.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=2000,
                messages=[{"role": "user", "content": [
                    {"type": "document", "source": {
                        "type": "base64", "media_type": "application/pdf",
                        "data": base64.b64encode(contenuto).decode()
                    }},
                    {"type": "text", "text": "Estrai tutte le informazioni rilevanti: "
                     "ragione sociale, sede, codice fiscale, soci, amministratori, "
                     "capitale sociale, oggetto sociale."}
                ]}]
            )
            testo_visura = resp_visura.content[0].text
        except:
            reader = PyPDF2.PdfReader(io.BytesIO(contenuto))
            testo_visura = "".join(p.extract_text() or "" for p in reader.pages)
        testi_supplementari["Visura Camerale"] = testo_visura
        documenti_binari[visura.name] = contenuto
        st.success("✅ Visura caricata")

    # ── PULSANTE GENERA ───────────────────────────────────────────────────────

    st.markdown("---")

    # Abilita il bottone solo se c'è almeno un bilancio pronto
    # oppure almeno un documento supplementare
    ha_contenuto = len(bilanci_pronti) > 0 or len(testi_supplementari) > 0

    if not ha_contenuto:
        st.info("Carica almeno un bilancio o documento per procedere.")

    if st.button("🚀 Genera Report", disabled=not ha_contenuto):

        lingua_prompt = "in inglese" if lingua == "English" else "in italiano"
        client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

        # Reset lista report per questa nuova generazione
        st.session_state["reports_generati"] = []

        # ── Caso A: uno o più bilanci caricati → report per ciascuno ─────────
        if bilanci_pronti:
            totale = len(bilanci_pronti)
            progress = st.progress(0, text="Preparazione...")

            for idx, item in enumerate(bilanci_pronti):
                nome = item["ragione_sociale"]
                anno = item["anno"]
                info = item["info"]
                mode = info.get("mode", "text")
                pagine_inviate = len(info.get("selected_pages", []))
                totale_pagine = info.get("total_pages", 0)

                if idx > 0:
                    time.sleep(5)  # pausa tra bilanci per evitare rate limit

                progress.progress(
                    idx / totale,
                    text=f"Analisi {idx+1}/{totale}: {nome} ({anno})..."
                )

                # Testo supplementare per questo report
                testo_suppl = ""
                for fonte, testo in testi_supplementari.items():
                    testo_suppl += f"\n\n--- {fonte} ---\n{testo[:5000]}"

                nota_estrazione = (
                    f"Il bilancio è scansionato — leggi i dati dalle immagini."
                    if mode == "scanned" else
                    f"Sono state selezionate {pagine_inviate} pagine rilevanti "
                    f"su {totale_pagine} totali del bilancio completo."
                )

                prompt_testo = f"""Sei un analista M&A e finance di uno studio legale internazionale.
Analizza il bilancio allegato di **{nome}** (esercizio {anno}) e, se presenti, i documenti supplementari.
{nota_estrazione}
Produci un report strutturato in JSON, con tutti i testi {lingua_prompt}.
{testo_suppl}

Rispondi SOLO con un oggetto JSON valido, senza backtick, senza testo aggiuntivo.

REGOLE CRITICHE PER I VALORI NUMERICI:
- Tutti i valori in "dati_finanziari" e "struttura_debito" devono essere NUMERI PURI (es. 1234567), mai stringhe
- Indica l'unità di misura usata nel bilancio nel campo "unita" (es. "migliaia di euro", "milioni di euro", "euro")
- Converti TUTTI i valori nella stessa unità dichiarata in "unita"
- I valori negativi (perdite, debito netto) devono avere il segno negativo (es. -5000)
- Se un valore non è presente nel documento usa null, MAI 0 o "N/D"
- Non inventare valori: meglio null che un numero sbagliato

{{
  "nome_azienda": "{nome}",
  "overview": "",
  "core_business": "",
  "mercati": "",
  "dati_finanziari": {{
    "unita": "migliaia di euro",
    "ricavi": null,
    "ebitda": null,
    "ebit": null,
    "utile_netto": null,
    "ammortamenti": null,
    "totale_attivo": null,
    "patrimonio_netto": null,
    "cassa": null,
    "anno_riferimento": "{anno}"
  }},
  "struttura_debito": {{
    "indebitamento_totale": null,
    "debito_bancario": null,
    "obbligazioni": null,
    "debito_netto": null,
    "leva_finanziaria": null,
    "scadenze_principali": "",
    "note": ""
  }},
  "ownership": {{
    "azionista_principale": "",
    "quota_principale": "",
    "altri_azionisti": "",
    "struttura_controllo": "",
    "note": ""
  }},
  "operazioni_ma": [
    {{"anno": "", "tipo": "", "descrizione": ""}}
  ],
  "note_aggiuntive": ""
}}

Se un testo non è disponibile scrivi N/D."""

                for tentativo in range(3):
                    try:
                        messaggio = client.messages.create(
                            model="claude-sonnet-4-6",
                            max_tokens=4000,
                            messages=[{"role": "user", "content": [
                                {
                                    "type": "document",
                                    "source": {
                                        "type": "base64",
                                        "media_type": "application/pdf",
                                        "data": item["pdf_b64"],
                                    }
                                },
                                {"type": "text", "text": prompt_testo}
                            ]}]
                        )
                        risposta = messaggio.content[0].text.strip()
                        if risposta.startswith("```"):
                            risposta = risposta.split("```")[1]
                            if risposta.startswith("json"):
                                risposta = risposta[4:]
                        report = json.loads(risposta.strip())

                        # Normalizza KPI a euro interi e valida
                        report = normalizza_kpi(report)
                        kpi_warnings = valida_kpi(report)
                        report["_validation_warnings"] = kpi_warnings

                        salva_report(nome, report, documenti_binari)

                        # Salva in session_state (lista per multi-report)
                        if "reports_generati" not in st.session_state:
                            st.session_state["reports_generati"] = []
                        st.session_state["reports_generati"].append({
                            "nome": nome,
                            "anno": anno,
                            "report": report,
                        })
                        # Compatibilità con il codice esistente
                        st.session_state["report"] = report
                        break

                    except Exception as e:
                        if tentativo < 2:
                            st.warning(f"{nome}: tentativo {tentativo+1} fallito, riprovo...")
                            time.sleep(30)
                        else:
                            st.error(f"❌ Errore per {nome}: {e}")

            progress.progress(1.0, text="Completato!")
            st.success(f"✅ {len(bilanci_pronti)} report generati e salvati!")

            # ── Pulsanti export ───────────────────────────────────────────────
            reports_list = [
                item["report"]
                for item in st.session_state.get("reports_generati", [])
                if item.get("report")
            ]
            if reports_list:
                nome_file = reports_list[0].get("nome_azienda", "report").replace(" ", "_")
                data_oggi_str = datetime.now().strftime("%Y%m%d")

                col_xl, col_wd = st.columns(2)

                with col_xl:
                    try:
                        excel_bytes = genera_excel(reports_list)
                        st.download_button(
                            label="📥 Scarica Excel",
                            data=excel_bytes,
                            file_name=f"TaxiReport_{nome_file}_{data_oggi_str}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True,
                        )
                    except Exception as e:
                        st.error(f"Errore Excel: {e}")

                with col_wd:
                    try:
                        word_bytes = genera_word(reports_list)
                        st.download_button(
                            label="📄 Scarica Word",
                            data=word_bytes,
                            file_name=f"TaxiReport_{nome_file}_{data_oggi_str}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                    except Exception as e:
                        st.error(f"Errore Word: {e}")

        # ── Caso B: solo documenti supplementari (comportamento originale) ────
        else:
            nome_azienda = st.text_input(
                "Nome dell'azienda", placeholder="es. Eco Eridania S.p.A.",
                key="nome_solo_suppl"
            )
            if not nome_azienda:
                st.warning("Inserisci il nome dell'azienda.")
            else:
                with st.spinner("Claude sta analizzando i documenti..."):
                    testo_completo = ""
                    for fonte, testo in testi_supplementari.items():
                        testo_completo += f"\n\n--- FONTE: {fonte} ---\n{testo[:10000]}"

                    lingua_prompt2 = "in inglese" if lingua == "English" else "in italiano"
                    prompt2 = f"""Sei un analista M&A e finance di uno studio legale internazionale.
Analizza i seguenti documenti relativi all'azienda {nome_azienda} e produci un report in JSON {lingua_prompt2}.

{testo_completo}

Rispondi SOLO con un oggetto JSON valido:

{{
  "nome_azienda": "",
  "overview": "",
  "core_business": "",
  "mercati": "",
  "dati_finanziari": {{
    "ricavi": "",
    "ebitda": "",
    "utile_netto": "",
    "totale_attivo": "",
    "patrimonio_netto": "",
    "anno_riferimento": ""
  }},
  "struttura_debito": {{
    "indebitamento_totale": "",
    "debito_bancario": "",
    "obbligazioni": "",
    "debito_netto": "",
    "leva_finanziaria": "",
    "scadenze_principali": "",
    "note": ""
  }},
  "ownership": {{
    "azionista_principale": "",
    "quota_principale": "",
    "altri_azionisti": "",
    "struttura_controllo": "",
    "note": ""
  }},
  "operazioni_ma": [
    {{"anno": "", "tipo": "", "descrizione": ""}}
  ],
  "note_aggiuntive": ""
}}

Se un dato non è disponibile scrivi N/D. Non inventare dati."""

                    try:
                        messaggio = client.messages.create(
                            model="claude-haiku-4-5-20251001",
                            max_tokens=4000,
                            messages=[{"role": "user", "content": prompt2}]
                        )
                        risposta = messaggio.content[0].text.strip()
                        if risposta.startswith("```"):
                            risposta = risposta.split("```")[1]
                            if risposta.startswith("json"):
                                risposta = risposta[4:]
                        report = json.loads(risposta.strip())
                        salva_report(nome_azienda, report, documenti_binari)
                        st.session_state["reports_generati"] = [{
                            "nome": nome_azienda,
                            "anno": "",
                            "report": report,
                        }]
                        st.success("✅ Report generato e salvato!")
                    except Exception as e:
                        st.error(f"Errore: {type(e).__name__}: {str(e)}")

    # ── VISUALIZZAZIONE REPORT ────────────────────────────────────────────────

    # Mostra tutti i report generati in questa sessione (multi-bilancio)
    if "reports_generati" in st.session_state and st.session_state["reports_generati"]:
        for item_r in st.session_state["reports_generati"]:
            report = item_r["report"]
            nome_r = item_r["nome"]
            anno_r = item_r["anno"]
            fin = report.get("dati_finanziari", {})

            st.markdown(f"## 📋 {nome_r} — {anno_r}")

            # Pannello validazione
            kpi_warnings = report.get("_validation_warnings", [])
            render_pannello_validazione(kpi_warnings)

            # Calcolo margine EBITDA
            r_val = fin.get("ricavi")
            e_val = fin.get("ebitda")
            margine_str = "N/D"
            if r_val and e_val and report.get("_normalizzato"):
                try:
                    margine_str = f"{float(e_val)/float(r_val)*100:.1f}%"
                except:
                    pass

            st.markdown(f"""
            <div class="kpi-grid">
                <div class="kpi-card"><div class="kpi-label">Ricavi</div><div class="kpi-value">{fmt_kpi(report,'ricavi')}</div></div>
                <div class="kpi-card"><div class="kpi-label">EBITDA</div><div class="kpi-value">{fmt_kpi(report,'ebitda')}</div></div>
                <div class="kpi-card"><div class="kpi-label">Margine EBITDA</div><div class="kpi-value">{margine_str}</div></div>
                <div class="kpi-card"><div class="kpi-label">Utile Netto</div><div class="kpi-value">{fmt_kpi(report,'utile_netto')}</div></div>
                <div class="kpi-card"><div class="kpi-label">Totale Attivo</div><div class="kpi-value">{fmt_kpi(report,'totale_attivo')}</div></div>
                <div class="kpi-card"><div class="kpi-label">Patrimonio Netto</div><div class="kpi-value">{fmt_kpi(report,'patrimonio_netto')}</div></div>
                <div class="kpi-card"><div class="kpi-label">Anno Rif.</div><div class="kpi-value">{fin.get('anno_riferimento','N/D')}</div></div>
            </div>
            """, unsafe_allow_html=True)

            with st.expander("🏢 Overview", expanded=True):
                st.markdown(f'<div class="section-box"><div class="section-title">Overview</div><div class="section-text">{report.get("overview","N/D")}</div></div>', unsafe_allow_html=True)

            with st.expander("⚙️ Core Business & Mercati"):
                st.markdown(f'<div class="section-box"><div class="section-title">Core Business</div><div class="section-text">{report.get("core_business","N/D")}</div></div>', unsafe_allow_html=True)
                st.markdown(f'<div class="section-box"><div class="section-title">Mercati</div><div class="section-text">{report.get("mercati","N/D")}</div></div>', unsafe_allow_html=True)

            with st.expander("💰 Struttura del Debito"):
                debito = report.get("struttura_debito", {})
                if isinstance(debito, dict):
                    st.markdown(f"""
                    <div class="kpi-grid">
                        <div class="kpi-card"><div class="kpi-label">Indebitamento Totale</div><div class="kpi-value">{fmt_kpi(report,'indebitamento_totale','struttura_debito')}</div></div>
                        <div class="kpi-card"><div class="kpi-label">Debito Bancario</div><div class="kpi-value">{fmt_kpi(report,'debito_bancario','struttura_debito')}</div></div>
                        <div class="kpi-card"><div class="kpi-label">Obbligazioni</div><div class="kpi-value">{fmt_kpi(report,'obbligazioni','struttura_debito')}</div></div>
                        <div class="kpi-card"><div class="kpi-label">Debito Netto</div><div class="kpi-value">{fmt_kpi(report,'debito_netto','struttura_debito')}</div></div>
                        <div class="kpi-card"><div class="kpi-label">Leva Finanziaria</div><div class="kpi-value">{debito.get('leva_finanziaria','N/D')}</div></div>
                    </div>
                    """, unsafe_allow_html=True)
                    if debito.get('scadenze_principali', 'N/D') != 'N/D':
                        st.markdown(f'<div class="section-box"><div class="section-title">Scadenze Principali</div><div class="section-text">{debito.get("scadenze_principali","N/D")}</div></div>', unsafe_allow_html=True)
                    if debito.get('note', 'N/D') != 'N/D':
                        st.markdown(f'<div class="section-box"><div class="section-title">Note</div><div class="section-text">{debito.get("note","N/D")}</div></div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="section-box"><div class="section-text">{debito}</div></div>', unsafe_allow_html=True)

            with st.expander("👥 Struttura Ownership"):
                ownership = report.get("ownership", {})
                if isinstance(ownership, dict):
                    st.markdown(f"""
                    <div class="kpi-grid">
                        <div class="kpi-card"><div class="kpi-label">Azionista Principale</div><div class="kpi-value">{ownership.get('azionista_principale','N/D')}</div></div>
                        <div class="kpi-card"><div class="kpi-label">Quota</div><div class="kpi-value">{ownership.get('quota_principale','N/D')}</div></div>
                    </div>
                    """, unsafe_allow_html=True)
                    if ownership.get('altri_azionisti', 'N/D') != 'N/D':
                        st.markdown(f'<div class="section-box"><div class="section-title">Altri Azionisti</div><div class="section-text">{ownership.get("altri_azionisti","N/D")}</div></div>', unsafe_allow_html=True)
                    if ownership.get('struttura_controllo', 'N/D') != 'N/D':
                        st.markdown(f'<div class="section-box"><div class="section-title">Struttura di Controllo</div><div class="section-text">{ownership.get("struttura_controllo","N/D")}</div></div>', unsafe_allow_html=True)
                    if ownership.get('note', 'N/D') != 'N/D':
                        st.markdown(f'<div class="section-box"><div class="section-title">Note</div><div class="section-text">{ownership.get("note","N/D")}</div></div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="section-box"><div class="section-text">{ownership}</div></div>', unsafe_allow_html=True)

            with st.expander("🔀 Operazioni M&A"):
                operazioni = report.get("operazioni_ma", [])
                if operazioni and operazioni[0].get("descrizione", "N/D") != "N/D":
                    for op in operazioni:
                        st.markdown(f"""
                        <div class="ma-item">
                            <span class="ma-anno">{op.get('anno','')}</span>
                            <div class="ma-tipo">{op.get('tipo','')}</div>
                            <div class="ma-desc">{op.get('descrizione','')}</div>
                        </div>""", unsafe_allow_html=True)
                else:
                    st.write("Nessuna operazione rilevata.")

            with st.expander("📝 Note Aggiuntive"):
                st.markdown(f'<div class="section-box"><div class="section-text">{report.get("note_aggiuntive","N/D")}</div></div>', unsafe_allow_html=True)

            # Confronto YoY se ci sono più report dello stesso soggetto
            reports_stesso_soggetto = [
                r for r in st.session_state.get("reports_generati", [])
                if r["nome"] == nome_r and r["anno"] != anno_r
            ]
            if reports_stesso_soggetto:
                yoy_warnings = []
                for r_prec in reports_stesso_soggetto:
                    yoy_warnings += confronta_anni(report, r_prec["report"])
                if yoy_warnings:
                    with st.expander("📊 Variazioni YoY anomale", expanded=True):
                        render_pannello_validazione(yoy_warnings)

            st.markdown("---")

    # ── Bottoni download (se ci sono report generati) ─────────────────────────
    if st.session_state.get("reports_generati"):
        reports_list = st.session_state["reports_generati"]
        st.markdown("### 💾 Esporta")
        col_xl, col_wd = st.columns(2)

        with col_xl:
            try:
                excel_bytes = genera_excel(reports_list)
                nome_file_xl = f"TaxiReport_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
                st.download_button(
                    label="📊 Scarica Excel",
                    data=excel_bytes,
                    file_name=nome_file_xl,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"Errore generazione Excel: {e}")

        with col_wd:
            try:
                word_bytes = genera_word(reports_list)
                nome_file_wd = f"TaxiReport_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                st.download_button(
                    label="📄 Scarica Word",
                    data=word_bytes,
                    file_name=nome_file_wd,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"Errore generazione Word: {e}")

        st.markdown("---")

    # Compatibilità con singolo report (aperto dall'archivio o caso B)
    elif "report" in st.session_state and "reports_generati" not in st.session_state:
        report = st.session_state["report"]
        # Normalizza se necessario (report vecchi pre-M2)
        if not report.get("_normalizzato"):
            report = normalizza_kpi(report)
            st.session_state["report"] = report
        fin    = report.get("dati_finanziari", {})
        debito = report.get("struttura_debito", {})

        st.markdown(f"## 📋 {report.get('nome_azienda', '')}")

        # Pannello validazione
        kpi_warnings = report.get("_validation_warnings") or valida_kpi(report)
        render_pannello_validazione(kpi_warnings)

        # Margine EBITDA
        r_val = fin.get("ricavi")
        e_val = fin.get("ebitda")
        margine_str = "N/D"
        if r_val and e_val:
            try:
                margine_str = f"{float(e_val)/float(r_val)*100:.1f}%"
            except:
                pass

        st.markdown(f"""
        <div class="kpi-grid">
            <div class="kpi-card"><div class="kpi-label">Ricavi</div><div class="kpi-value">{fmt_kpi(report,'ricavi')}</div></div>
            <div class="kpi-card"><div class="kpi-label">EBITDA</div><div class="kpi-value">{fmt_kpi(report,'ebitda')}</div></div>
            <div class="kpi-card"><div class="kpi-label">Margine EBITDA</div><div class="kpi-value">{margine_str}</div></div>
            <div class="kpi-card"><div class="kpi-label">Utile Netto</div><div class="kpi-value">{fmt_kpi(report,'utile_netto')}</div></div>
            <div class="kpi-card"><div class="kpi-label">Totale Attivo</div><div class="kpi-value">{fmt_kpi(report,'totale_attivo')}</div></div>
            <div class="kpi-card"><div class="kpi-label">Patrimonio Netto</div><div class="kpi-value">{fmt_kpi(report,'patrimonio_netto')}</div></div>
            <div class="kpi-card"><div class="kpi-label">Anno Rif.</div><div class="kpi-value">{fin.get('anno_riferimento','N/D')}</div></div>
        </div>
        """, unsafe_allow_html=True)

        with st.expander("🏢 Overview", expanded=True):
            st.markdown(f'<div class="section-box"><div class="section-title">Overview</div><div class="section-text">{report.get("overview","N/D")}</div></div>', unsafe_allow_html=True)

        with st.expander("⚙️ Core Business & Mercati"):
            st.markdown(f'<div class="section-box"><div class="section-title">Core Business</div><div class="section-text">{report.get("core_business","N/D")}</div></div>', unsafe_allow_html=True)
            st.markdown(f'<div class="section-box"><div class="section-title">Mercati</div><div class="section-text">{report.get("mercati","N/D")}</div></div>', unsafe_allow_html=True)

        with st.expander("💰 Struttura del Debito"):
            if isinstance(debito, dict):
                st.markdown(f"""
                <div class="kpi-grid">
                    <div class="kpi-card"><div class="kpi-label">Indebitamento Totale</div><div class="kpi-value">{fmt_kpi(report,'indebitamento_totale','struttura_debito')}</div></div>
                    <div class="kpi-card"><div class="kpi-label">Debito Bancario</div><div class="kpi-value">{fmt_kpi(report,'debito_bancario','struttura_debito')}</div></div>
                    <div class="kpi-card"><div class="kpi-label">Obbligazioni</div><div class="kpi-value">{fmt_kpi(report,'obbligazioni','struttura_debito')}</div></div>
                    <div class="kpi-card"><div class="kpi-label">Debito Netto</div><div class="kpi-value">{fmt_kpi(report,'debito_netto','struttura_debito')}</div></div>
                    <div class="kpi-card"><div class="kpi-label">Leva Finanziaria</div><div class="kpi-value">{debito.get('leva_finanziaria','N/D')}</div></div>
                </div>
                """, unsafe_allow_html=True)
                if debito.get('scadenze_principali', 'N/D') not in ('N/D', '', None):
                    st.markdown(f'<div class="section-box"><div class="section-title">Scadenze Principali</div><div class="section-text">{debito.get("scadenze_principali")}</div></div>', unsafe_allow_html=True)
                if debito.get('note', 'N/D') not in ('N/D', '', None):
                    st.markdown(f'<div class="section-box"><div class="section-title">Note</div><div class="section-text">{debito.get("note")}</div></div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="section-box"><div class="section-text">{debito}</div></div>', unsafe_allow_html=True)

        with st.expander("👥 Struttura Ownership"):
            ownership = report.get("ownership", {})
            if isinstance(ownership, dict):
                st.markdown(f"""
                <div class="kpi-grid">
                    <div class="kpi-card"><div class="kpi-label">Azionista Principale</div><div class="kpi-value">{ownership.get('azionista_principale','N/D')}</div></div>
                    <div class="kpi-card"><div class="kpi-label">Quota</div><div class="kpi-value">{ownership.get('quota_principale','N/D')}</div></div>
                </div>
                """, unsafe_allow_html=True)
                for lbl, key in [("Altri Azionisti","altri_azionisti"),
                                  ("Struttura di Controllo","struttura_controllo"),
                                  ("Note","note")]:
                    val = ownership.get(key, "N/D")
                    if val not in ("N/D", "", None):
                        st.markdown(f'<div class="section-box"><div class="section-title">{lbl}</div><div class="section-text">{val}</div></div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="section-box"><div class="section-text">{ownership}</div></div>', unsafe_allow_html=True)

        with st.expander("🔀 Operazioni M&A"):
            operazioni = report.get("operazioni_ma", [])
            if operazioni and operazioni[0].get("descrizione","N/D") != "N/D":
                for op in operazioni:
                    st.markdown(f"""
                    <div class="ma-item">
                        <span class="ma-anno">{op.get('anno','')}</span>
                        <div class="ma-tipo">{op.get('tipo','')}</div>
                        <div class="ma-desc">{op.get('descrizione','')}</div>
                    </div>""", unsafe_allow_html=True)
            else:
                st.write("Nessuna operazione rilevata.")

        with st.expander("📝 Note Aggiuntive"):
            st.markdown(f'<div class="section-box"><div class="section-text">{report.get("note_aggiuntive","N/D")}</div></div>', unsafe_allow_html=True)

        # Export anche per report singolo
        st.markdown("---")
        st.markdown("### 💾 Esporta")
        _exp_item = [{"nome": report.get("nome_azienda","Report"),
                      "anno": fin.get("anno_riferimento",""),
                      "report": report}]
        col_xl2, col_wd2 = st.columns(2)
        with col_xl2:
            try:
                st.download_button("📊 Scarica Excel", genera_excel(_exp_item),
                    file_name=f"TaxiReport_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True)
            except Exception as e:
                st.error(f"Errore Excel: {e}")
        with col_wd2:
            try:
                st.download_button("📄 Scarica Word", genera_word(_exp_item),
                    file_name=f"TaxiReport_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
            except Exception as e:
                st.error(f"Errore Word: {e}")

# ══════════════════════════════════════════════════════════════════════════════


# ══════════════════════════════════════════════════════════════════════════════
# PAGINA: ARCHIVIO
# ══════════════════════════════════════════════════════════════════════════════

elif st.session_state["pagina"] == "archivio":

    tab_archivio, tab_confronta = st.tabs(["🗂️ Archivio", "📊 Confronta aziende"])

    # ── TAB ARCHIVIO ──────────────────────────────────────────────────────────
    with tab_archivio:
        reports_arch = carica_archivio()

        if not reports_arch:
            st.info("Nessun report salvato. Genera il tuo primo report!")
        else:
            # Ricerca
            cerca = st.text_input("🔍 Cerca per nome azienda", placeholder="es. Angelini...",
                                  label_visibility="collapsed")
            if cerca:
                reports_arch = [x for x in reports_arch
                                if cerca.lower() in x.get("nome","").lower()
                                or cerca.lower() in x["report"].get("nome_azienda","").lower()]

            st.caption(f"{len(reports_arch)} report trovati")

            for item in reports_arch:
                r   = item["report"]
                fin = r.get("dati_finanziari", {})
                nome_az = r.get("nome_azienda", item.get("nome", "—"))
                anno_az = fin.get("anno_riferimento", "—")
                warnings_count = len(r.get("_validation_warnings", []))
                warn_badge = (f' <span style="color:#e65100;font-size:11px;">'
                              f'⚠️ {warnings_count} warning</span>'
                              if warnings_count else "")

                # Report con schema flessibile (caricati da docx)
                is_flex = "sezioni" in r

                if is_flex:
                    # Anteprima: mostra primi campi corti dalle prime sezioni
                    kpi_items_html = ""
                    for sez in r.get("sezioni", [])[:3]:
                        for campo in sez.get("campi", [])[:2]:
                            lbl = campo.get("label", "")
                            val = str(campo.get("valore", ""))[:40]
                            if len(val) <= 40:
                                kpi_items_html += f"""
                                <div class="archivio-kpi-item">
                                    <span class="archivio-kpi-label">{lbl}</span>
                                    <span class="archivio-kpi-val">{val}</span>
                                </div>"""
                    num_sezioni = len(r.get("sezioni", []))
                    st.markdown(f"""
                    <div class="archivio-card">
                        <div class="archivio-nome">{nome_az}{warn_badge}</div>
                        <div class="archivio-data">📅 {item["data"]} &nbsp;|&nbsp; 📄 Report caricato ({num_sezioni} sezioni)</div>
                        <div class="archivio-kpi">{kpi_items_html}</div>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f"""
                    <div class="archivio-card">
                        <div class="archivio-nome">{nome_az}{warn_badge}</div>
                        <div class="archivio-data">📅 {item["data"]} &nbsp;|&nbsp; Anno bilancio: {anno_az}</div>
                        <div class="archivio-kpi">
                            <div class="archivio-kpi-item">
                                <span class="archivio-kpi-label">Ricavi</span>
                                <span class="archivio-kpi-val">{_fmt(fin.get("ricavi"))}</span>
                            </div>
                            <div class="archivio-kpi-item">
                                <span class="archivio-kpi-label">EBITDA</span>
                                <span class="archivio-kpi-val">{_fmt(fin.get("ebitda"))}</span>
                            </div>
                            <div class="archivio-kpi-item">
                                <span class="archivio-kpi-label">Utile Netto</span>
                                <span class="archivio-kpi-val">{_fmt(fin.get("utile_netto"))}</span>
                            </div>
                            <div class="archivio-kpi-item">
                                <span class="archivio-kpi-label">PN</span>
                                <span class="archivio-kpi-val">{_fmt(fin.get("patrimonio_netto"))}</span>
                            </div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)

                col_a, col_b, col_c = st.columns([1, 1, 4])
                with col_a:
                    if st.button("📖 Apri", key=f"apri_{item['riga']}"):
                        if is_flex:
                            # Report con schema flessibile → apri in carica_esistente
                            st.session_state["report_caricato"] = r
                            st.session_state["report_caricato_nome"] = nome_az
                            st.session_state["report_caricato_file"] = ""
                            st.session_state.pop("reports_generati", None)
                            st.session_state["pagina"] = "carica_esistente"
                        else:
                            # Report con schema classico → apri in genera
                            st.session_state["report"] = r
                            st.session_state.pop("reports_generati", None)
                            st.session_state["pagina"] = "genera"
                        st.rerun()
                with col_b:
                    if st.button("🗑️ Elimina", key=f"elimina_{item['riga']}"):
                        try:
                            sheet = get_sheet()
                            sheet.delete_rows(item["riga"])
                            st.success("Report eliminato.")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Errore: {e}")

    # ── TAB CONFRONTA ─────────────────────────────────────────────────────────
    with tab_confronta:
        reports_tutti = carica_archivio()

        if len(reports_tutti) < 2:
            st.info("Servono almeno 2 report in archivio per confrontare.")
        else:
            # Selezione aziende
            opzioni = [
                f"{x['report'].get('nome_azienda', x.get('nome','—'))} "
                f"({x['report'].get('dati_finanziari',{}).get('anno_riferimento','—')})"
                for x in reports_tutti
            ]
            selezionati = st.multiselect(
                "Seleziona 2-5 aziende/anni da confrontare",
                options=opzioni,
                max_selections=5,
            )

            # Benchmark settore opzionale
            settore_sel = st.selectbox(
                "Benchmark di settore (opzionale)",
                options=["— Nessun benchmark —"] + list(BENCHMARK_SETTORE.keys()),
            )
            bm = BENCHMARK_SETTORE.get(settore_sel) if settore_sel != "— Nessun benchmark —" else None

            if len(selezionati) < 2:
                st.caption("Seleziona almeno 2 aziende per visualizzare il confronto.")
            else:
                # Recupera i report selezionati
                idx_sel = [opzioni.index(s) for s in selezionati]
                sel_reports = [reports_tutti[i] for i in idx_sel]

                st.markdown("---")
                st.markdown("#### Tabella comparativa")

                # Righe KPI da confrontare
                kpi_conf = [
                    ("Ricavi",            "dati_finanziari",  "ricavi",             None),
                    ("EBITDA",            "dati_finanziari",  "ebitda",             None),
                    ("Margine EBITDA",    None,               "_margine_ebitda",    "margine_ebitda"),
                    ("EBIT",              "dati_finanziari",  "ebit",               None),
                    ("Utile Netto",       "dati_finanziari",  "utile_netto",        None),
                    ("Margine Utile",     None,               "_margine_utile",     "margine_utile"),
                    ("Totale Attivo",     "dati_finanziari",  "totale_attivo",      None),
                    ("Patrimonio Netto",  "dati_finanziari",  "patrimonio_netto",   None),
                    ("Debito Netto",      "struttura_debito", "debito_netto",       None),
                    ("Leva (x)",          "struttura_debito", "leva_finanziaria",   "leva"),
                ]

                # Intestazione tabella
                header_cols = st.columns([2] + [2] * len(sel_reports) + ([1] if bm else []))
                header_cols[0].markdown("**KPI**")
                for j, s in enumerate(selezionati):
                    header_cols[j+1].markdown(f"**{s}**")
                if bm:
                    header_cols[-1].markdown("**Benchmark**")

                st.markdown('<hr style="border-color:#444;margin:4px 0 8px 0;">', unsafe_allow_html=True)

                for label, sezione, campo, bm_key in kpi_conf:
                    row_cols = st.columns([2] + [2] * len(sel_reports) + ([1] if bm else []))
                    row_cols[0].markdown(f"<span style='color:#999;font-size:13px;'>{label}</span>",
                                         unsafe_allow_html=True)

                    vals = []
                    for j, item in enumerate(sel_reports):
                        r = item["report"]
                        fin_c = r.get("dati_finanziari", {})
                        deb_c = r.get("struttura_debito", {})
                        val = None

                        if campo == "_margine_ebitda":
                            ric = fin_c.get("ricavi")
                            ebt = fin_c.get("ebitda")
                            if ric and ebt:
                                try: val = float(ebt) / float(ric)
                                except: pass
                            display = f"{val*100:.1f}%" if val is not None else "N/D"
                        elif campo == "_margine_utile":
                            ric = fin_c.get("ricavi")
                            utn = fin_c.get("utile_netto")
                            if ric and utn:
                                try: val = float(utn) / float(ric)
                                except: pass
                            display = f"{val*100:.1f}%" if val is not None else "N/D"
                        elif sezione == "dati_finanziari":
                            raw = fin_c.get(campo)
                            if raw is not None:
                                try: val = float(raw)
                                except: pass
                            display = _fmt(val)
                        elif sezione == "struttura_debito":
                            raw = deb_c.get(campo)
                            if raw is not None:
                                try: val = float(raw)
                                except: pass
                            display = f"{val:.1f}x" if (campo == "leva_finanziaria" and val) else _fmt(val)
                        else:
                            display = "N/D"

                        vals.append(val)
                        row_cols[j+1].markdown(
                            f"<span style='font-size:13px;font-weight:600;"
                            f"color:#c8e04a;'>{display}</span>",
                            unsafe_allow_html=True
                        )

                    # Colonna benchmark
                    if bm and bm_key and bm_key in bm:
                        # Usa il valore medio dei selezionati per confronto
                        valori_validi = [v for v in vals if v is not None]
                        val_medio = sum(valori_validi)/len(valori_validi) if valori_validi else None
                        testo_bm, colore_bm = _benchmark_cell(val_medio, bm.get(bm_key))
                        lo, hi = bm[bm_key]
                        range_str = (f"{lo*100:.0f}–{hi*100:.0f}%"
                                     if "margine" in bm_key
                                     else f"{lo:.1f}–{hi:.1f}x")
                        row_cols[-1].markdown(
                            f"<span style='font-size:11px;color:{colore_bm};'>"
                            f"{testo_bm}<br>"
                            f"<span style='color:#666;'>Range: {range_str}</span>"
                            f"</span>",
                            unsafe_allow_html=True
                        )
                    elif bm:
                        row_cols[-1].markdown(
                            "<span style='font-size:11px;color:#666;'>—</span>",
                            unsafe_allow_html=True
                        )

                    st.markdown('<hr style="border-color:#2a2a2a;margin:2px 0;">', unsafe_allow_html=True)

                # Export comparazione
                st.markdown("---")
                try:
                    reports_export = [
                        {"nome": x["report"].get("nome_azienda", x.get("nome","—")),
                         "anno": x["report"].get("dati_finanziari",{}).get("anno_riferimento","—"),
                         "report": x["report"]}
                        for x in sel_reports
                    ]
                    excel_bytes = genera_excel(reports_export)
                    st.download_button(
                        label="📊 Scarica confronto Excel",
                        data=excel_bytes,
                        file_name=f"Confronto_{datetime.now().strftime('%Y%m%d')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True,
                    )
                except Exception as e:
                    st.error(f"Errore export: {e}")

# ══════════════════════════════════════════════════════════════════════════════
# PAGINA: CARICA REPORT ESISTENTE
# ══════════════════════════════════════════════════════════════════════════════

elif st.session_state["pagina"] == "carica_esistente":

    st.markdown("### 📄 Carica Report Esistente")
    st.caption(
        "Carica un file Word (.docx) di un taxi report già compilato. "
        "Il sistema estrarrà e riprodurrà fedelmente tutte le sezioni del documento."
    )

    docx_file = st.file_uploader(
        "Seleziona un file Word (.docx) di un taxi report esistente",
        type=["docx"],
        accept_multiple_files=False,
        key="carica_report_docx",
        label_visibility="collapsed",
    )

    if docx_file:
        # ── Estrazione testo strutturato dal .docx ────────────────────────
        try:
            doc_bytes = docx_file.read()
            doc = Document(io.BytesIO(doc_bytes))

            # Estrai paragrafi
            paragrafi = [p.text for p in doc.paragraphs if p.text.strip()]

            # Estrai tabelle preservando struttura label → valore
            # Deduplica sia celle merged orizzontali (stessa riga)
            # sia celle merged verticali (stesse righe consecutive)
            testo_tabelle = []
            righe_viste_globale = set()
            for table in doc.tables:
                for row in table.rows:
                    celle_uniche = []
                    viste = set()
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t and t not in viste:
                            celle_uniche.append(t)
                            viste.add(t)
                    if celle_uniche:
                        riga_str = " | ".join(celle_uniche)
                        # Evita righe identiche da merge verticali
                        if riga_str not in righe_viste_globale:
                            righe_viste_globale.add(riga_str)
                            testo_tabelle.append(riga_str)

            testo_completo = "\n".join(paragrafi)
            if testo_tabelle:
                testo_completo += "\n\n--- TABELLE ---\n" + "\n".join(testo_tabelle)

        except Exception as e:
            testo_completo = None
            st.markdown(f"""
            <div class="section-box" style="border-left-color: #ef5350;">
                <div class="section-title" style="color: #ef5350;">
                    ❌ Errore lettura file
                </div>
                <div class="section-text">
                    Impossibile leggere il file Word: {type(e).__name__}: {e}<br>
                    Assicurati che il file sia un documento .docx valido e non corrotto.
                </div>
            </div>
            """, unsafe_allow_html=True)

        if testo_completo:
            with st.expander("📃 Anteprima testo estratto dal documento"):
                st.text(testo_completo[:5000] + ("..." if len(testo_completo) > 5000 else ""))
            st.caption(f"Estratti {len(paragrafi)} paragrafi e {len(doc.tables)} tabelle — {len(testo_completo)} caratteri totali")

            if st.button("🔍 Estrai dati dal report", use_container_width=True):
                client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY"))

                prompt_estrazione = f"""You are a legal M&A analyst. You receive the full text extracted from a BD (Business Development) taxi report document.

Your task: extract EVERY piece of information from this document and reproduce it faithfully in a structured JSON.

--- DOCUMENT TEXT ---
{testo_completo[:20000]}
--- END ---

CRITICAL RULES:
1. DO NOT invent, summarize, or paraphrase. Copy the content exactly as written in the document.
2. DO NOT skip any section, field, or piece of information. If it's in the document, it must be in the JSON.
3. Preserve the ORIGINAL LANGUAGE of the document (English, Italian, etc.) — do not translate.
4. For each section you find, create an entry in the "sezioni" array with the exact section title and ALL its fields.
5. Each field has a "label" (the field name as it appears) and a "valore" (the full content, verbatim).
6. If a field contains multiple paragraphs or bullet points, include them all in "valore" separated by newlines.
7. NEVER repeat the same content in multiple fields. Each piece of information must appear EXACTLY ONCE. If the same text appears under multiple labels in the source (due to merged cells), include it only in the most specific/relevant field.
8. If a cell contains multiple sub-sections (e.g. "Rumored Transactions" and "Critical Issues / Outlook" in the same cell), split them into SEPARATE fields with their own label and value.

Respond ONLY with valid JSON, no backticks, no additional text.

{{
  "nome_azienda": "",
  "sezioni": [
    {{
      "titolo_sezione": "General Information",
      "campi": [
        {{"label": "Company name", "valore": "..."}},
        {{"label": "Business Description", "valore": "... full text ..."}},
        {{"label": "Headquarter", "valore": "..."}},
        {{"label": "...", "valore": "..."}}
      ]
    }},
    {{
      "titolo_sezione": "Financial Data",
      "campi": [
        {{"label": "...", "valore": "..."}},
        {{"label": "...", "valore": "..."}}
      ]
    }}
  ]
}}

Extract ALL sections you find. Common sections include (but are not limited to):
General Information, Financial Data, Deals/Transactions, Governance, Legal Advisors, etc.
If you find sections with different names, use those names exactly."""

                with st.spinner("Claude sta analizzando il documento..."):
                    for tentativo in range(3):
                        try:
                            messaggio = client.messages.create(
                                model="claude-sonnet-4-20250514",
                                max_tokens=8000,
                                messages=[{"role": "user", "content": prompt_estrazione}]
                            )
                            risposta = messaggio.content[0].text.strip()
                            if risposta.startswith("```"):
                                risposta = risposta.split("```")[1]
                                if risposta.startswith("json"):
                                    risposta = risposta[4:]
                            report = json.loads(risposta.strip())

                            nome_azienda = report.get("nome_azienda", docx_file.name)
                            st.session_state["report_caricato"] = report
                            st.session_state["report_caricato_nome"] = nome_azienda
                            st.session_state["report_caricato_file"] = docx_file.name
                            st.success("✅ Estrazione completata!")
                            break

                        except json.JSONDecodeError as e:
                            if tentativo < 2:
                                st.warning(f"Tentativo {tentativo+1}: risposta non valida, riprovo...")
                                time.sleep(5)
                            else:
                                st.markdown(f"""
                                <div class="section-box" style="border-left-color: #ef5350;">
                                    <div class="section-title" style="color: #ef5350;">
                                        ❌ Errore di estrazione
                                    </div>
                                    <div class="section-text">
                                        Impossibile interpretare la risposta AI dopo 3 tentativi.<br>
                                        Dettaglio: {e}
                                    </div>
                                </div>
                                """, unsafe_allow_html=True)
                        except Exception as e:
                            if tentativo < 2:
                                st.warning(f"Tentativo {tentativo+1} fallito: {e}, riprovo...")
                                time.sleep(10)
                            else:
                                st.markdown(f"""
                                <div class="section-box" style="border-left-color: #ef5350;">
                                    <div class="section-title" style="color: #ef5350;">
                                        ❌ Errore
                                    </div>
                                    <div class="section-text">
                                        {type(e).__name__}: {e}
                                    </div>
                                </div>
                                """, unsafe_allow_html=True)

    # ── VISUALIZZAZIONE REPORT ESTRATTO DA DOCX ───────────────────────────

    if st.session_state.get("report_caricato"):
        report = st.session_state["report_caricato"]
        nome_az = st.session_state.get("report_caricato_nome", "")
        file_origine = st.session_state.get("report_caricato_file", "")

        st.markdown("---")
        st.markdown(f"## 📋 {nome_az}")
        if file_origine:
            st.caption(f"Estratto da: {file_origine}")

        sezioni = report.get("sezioni", [])

        if not sezioni:
            st.warning("Nessuna sezione estratta dal documento.")
        else:
            # Icone per le sezioni comuni
            ICONE_SEZIONI = {
                "general information": "🏢",
                "informazioni generali": "🏢",
                "financial data": "💰",
                "dati finanziari": "💰",
                "financial data of the group": "💰",
                "deals": "🔀",
                "transactions": "🔀",
                "recent transactions": "🔀",
                "deals/transactions": "🔀",
                "governance": "👥",
                "ownership": "👥",
                "legal advisors": "⚖️",
                "current preferred legal advisor": "⚖️",
                "critical issues": "⚠️",
                "critical issues / outlook": "⚠️",
                "outlook": "📊",
                "real estate portfolio": "🏗️",
                "rumored transactions": "📰",
            }

            for idx_s, sezione in enumerate(sezioni):
                titolo = sezione.get("titolo_sezione", f"Sezione {idx_s+1}")
                campi = sezione.get("campi", [])
                icona = ICONE_SEZIONI.get(titolo.lower().strip(), "📌")

                with st.expander(f"{icona} {titolo}", expanded=(idx_s == 0)):
                    if not campi:
                        st.write("Nessun dato in questa sezione.")
                        continue

                    # Identifica campi "corti" (KPI-like) vs "lunghi" (testo)
                    campi_corti = []
                    campi_lunghi = []
                    for c in campi:
                        val = str(c.get("valore", ""))
                        if len(val) <= 80 and "\n" not in val:
                            campi_corti.append(c)
                        else:
                            campi_lunghi.append(c)

                    # Mostra campi corti come KPI cards
                    if campi_corti:
                        cards_html = ""
                        for c in campi_corti:
                            lbl = c.get("label", "")
                            val = str(c.get("valore", "N/D"))
                            cards_html += f"""
                            <div class="kpi-card">
                                <div class="kpi-label">{lbl}</div>
                                <div class="kpi-value" style="font-size:14px;">{val}</div>
                            </div>"""
                        st.markdown(f'<div class="kpi-grid">{cards_html}</div>',
                                    unsafe_allow_html=True)

                    # Mostra campi lunghi come section-box
                    for c in campi_lunghi:
                        lbl = c.get("label", "")
                        val = str(c.get("valore", ""))
                        # Converti newline in <br> per HTML
                        val_html = val.replace("\n", "<br>")
                        st.markdown(f"""
                        <div class="section-box">
                            <div class="section-title">{lbl}</div>
                            <div class="section-text">{val_html}</div>
                        </div>""", unsafe_allow_html=True)

        # ── SALVATAGGIO ───────────────────────────────────────────────────

        st.markdown("---")

        if st.button("💾 Salva su Google Sheets", use_container_width=True, key="salva_caricato"):
            salva_report(nome_az, report, {})
            st.success("✅ Report salvato nell'archivio Google Sheets!")